# The Python-Docx third-party module can read and write .docx Word files.
# Open a Word file with docx.Document()
# Access one of the Paragraph objects from the "paragraphs" member variable, which is a list of Paragraph objects.
# Paragraph objects have a "text" member variable containing the text as a string value.
# Paragraphs are composed of "runs".  The "runs" member variable of a Paragraph object contains a list of Run objects.
# Run objects also have a "text" member variable.
# Run objects have a "bold", "italic", and "underline" member variables which can be set to True or False.
# Paragraph and run objects have a "style" member variable that can be set to one of Word's built-in styles.
# Word files can be created by calling add_paragraph() and add_run() to append text content.

# Import relevant Modules
import docx, os

# Update Working Directory to Local Lesson Folder where the example.xlsx file resides
os.chdir('C:\\Development\\Python-Projects\\Automate-The-Boring-Stuff-With-Python\\Section_14_Excel_Word_PDF_Documents\\Lesson_45_ReadingEditingWordDocs')

# Verify the Current Working Directory
filePath = os.getcwd()

# DEBUG Print the Current Working Directories File Path
print(filePath)

# Open a Docx Document and store it in a Document Object variable
docObject = docx.Document('demo.docx')

# Print the List of paragrpha objects and the associated text
print(docObject.paragraphs)                        # Print all the Paragraph Objects
print(docObject.paragraphs[0])                     # Print the 1st Paragraph Object at Index 0
print(docObject.paragraphs[0].text)                # Print the 1st Paragraph Object at Index 0 text
print(docObject.paragraphs[1].text)                # Print the 1st Paragraph Object at Index 1 text

# Create a new Paragraph Object Variable
paragraphObject1 = docObject.paragraphs[1]

# Print the Paragraph Objects Runs values and their associated properties
print(paragraphObject1.runs)                        # Print all the Paragraph Objects Runs values
print(paragraphObject1.runs[0].text)                # Print all the Paragraph Objects Runs value for Index 0
print(paragraphObject1.runs[0].bold)                # Print all the Paragraph Objects Runs value for Index 0
print(paragraphObject1.runs[0].italic)              # Print all the Paragraph Objects Runs value for Index 0
print(paragraphObject1.runs[0].underline)           # Print all the Paragraph Objects Runs value for Index 0
print(paragraphObject1.runs[1].text)                # Print all the Paragraph Objects Runs value for Index 1
print(paragraphObject1.runs[1].bold)                # Print all the Paragraph Objects Runs value for Index 1
print(paragraphObject1.runs[1].italic)              # Print all the Paragraph Objects Runs value for Index 1
print(paragraphObject1.runs[1].underline)           # Print all the Paragraph Objects Runs value for Index 1
print(paragraphObject1.runs[2].text)                # Print all the Paragraph Objects Runs value for Index 2
print(paragraphObject1.runs[2].bold)                # Print all the Paragraph Objects Runs value for Index 2
print(paragraphObject1.runs[2].italic)              # Print all the Paragraph Objects Runs value for Index 2
print(paragraphObject1.runs[2].underline)           # Print all the Paragraph Objects Runs value for Index 2
print(paragraphObject1.runs[3].text)                # Print all the Paragraph Objects Runs value for Index 3
print(paragraphObject1.runs[3].bold)                # Print all the Paragraph Objects Runs value for Index 3
print(paragraphObject1.runs[3].italic)              # Print all the Paragraph Objects Runs value for Index 3
print(paragraphObject1.runs[3].underline)           # Print all the Paragraph Objects Runs value for Index 3

# Update the Paragraphs Text properties
paragraphObject1.runs[3].underline = True
paragraphObject1.runs[3].text = 'italic and underlined.'

# Save the Document changes
docObject.save('demo2.docx')

# Print the Paragraph Objects Style
print(paragraphObject1.style)

# Update the Paragraph Objects Style
paragraphObject1.style = 'Title'

# Save the Document changes again
docObject.save('demo3.docx')