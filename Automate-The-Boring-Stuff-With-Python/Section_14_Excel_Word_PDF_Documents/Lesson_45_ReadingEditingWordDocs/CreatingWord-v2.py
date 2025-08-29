# The Python-Docx third-party module can read and write .docx Word files.
# Open a Word file with docx.Document()
# Access one of the Paragraph objects from the "paragraphs" member variable, which is a list of Paragraph objects.
# Paragraph objects have a "text" member variable containing the text as a string value.
# Paragraphs are composed of "runs".  The "runs" member variable of a Paragraph object contains a list of Run objects.
# Run objects also have a "text" member variable.
# Run objects have a "bold", "italic", and "underline" member variables which can be set to True or False.
# Paragraph and run objects have a "style" member variable that can be set to one of Word's built-in styles.
# Word files can be created by calling add_paragraph() and add_run() to append text content.

# Import relevant Modules
import docx, os

# Update Working Directory to Local Lesson Folder where the example.xlsx file resides
os.chdir('C:\\Development\\Python-Projects\\Automate-The-Boring-Stuff-With-Python\\Section_14_Excel_Word_PDF_Documents\\Lesson_45_ReadingEditingWordDocs')

# Verify the Current Working Directory
filePath = os.getcwd()

# DEBUG Print the Current Working Directories File Path
print(filePath)

# Create a new Docx Document Object variable
docObject = docx.Document()

# Create a new Paragraph in the Docx Object
docObject.add_paragraph('Hello this is a paragraph.')
docObject.add_paragraph('This is another paragraph.')

# Save the Document changes again
docObject.save('demo4.docx')

# Create a new Paragraph Variable
paragraphObject1 = docObject.paragraphs[0]

# Create a new Run for the Paragraph Variable
paragraphObject1.add_run('This is a new run.')

# Print the present run value
print(paragraphObject1.runs)

# Update the Runs value to Bold at Index 1
paragraphObject1.runs[1].bold = True

# Save the Document changes again
docObject.save('demo5.docx')