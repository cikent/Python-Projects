# The Python-Docx third-party module can read and write .docx Word files.
# Open a Word file with docx.Document()
# Access one of the Paragraph objects from the "paragraphs" member variable, which is a list of Paragraph objects.
# Paragraph objects have a "text" member variable containing the text as a string value.
# Paragraphs are composed of "runs".  The "runs" member variable of a Paragraph object contains a list of Run objects.
# Run objects also have a "text" member variable.
# Run objects have a "bold", "italic", and "underline" member variables which can be set to True or False.
# Paragraph and run objects have a "style" member variable that can be set to one of Word's built-in styles.
# Word files can be created by calling add_paragraph() and add_run() to append text content.

# Import relevant Modules
import docx, os

# Update Working Directory to Local Lesson Folder where the example.xlsx file resides
os.chdir('C:\\Development\\Python-Projects\\Automate-The-Boring-Stuff-With-Python\\Section_14_Excel_Word_PDF_Documents\\Lesson_45_ReadingEditingWordDocs')

# Verify the Current Working Directory
filePath = os.getcwd()

# DEBUG Print the Current Working Directories File Path
print(filePath)

# Create a Function to grab all of the text inside a word Document as a String then Print it
def getText(filename):
    doc = docx.Document(filename)
    fullText = []                   # Create an empty List
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

print(getText('demo.docx'))